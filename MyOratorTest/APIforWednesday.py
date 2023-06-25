import docx
import chardet
from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.by import By
from selenium.webdriver.common.action_chains import ActionChains
import time
import pygame
import threading

# 全局变量，用于保存当前正在播放的音乐文件名
current_playing = None

def create_word_file(content, filename):
    # 创建一个Document对象
    document = docx.Document()
    
    # 添加段落
    paragraph = document.add_paragraph(content)
    
    # 保存文档
    document.save(filename)

def read_txt_file(filename):
    with open(filename, 'rb') as file:
        rawdata = file.read()
        result = chardet.detect(rawdata)
        encoding = result['encoding']
    with open(filename, 'r', encoding=encoding) as file:
        content = file.read()
    return content

def read_word_file(filename):
    # 打开 Word 文档
    doc = docx.Document(filename)
    
    # 读取文本内容
    content = []
    for para in doc.paragraphs:
        content.append(para.text)
    
    # 将内容转换为字符串并返回
    return "\n".join(content)


def open_website(url):
    # 设置浏览器驱动路径
    driver_path = "C:\Program Files\Google\Driver\chromedriver"

    options = webdriver.ChromeOptions()
    options.add_experimental_option('detach', True)  #不自动关闭浏览器
    options.add_argument('--start-maximized')#浏览器窗口最大化
    driver=webdriver.Chrome(options=options ,executable_path=driver_path)
    driver.get(url)
    return

def play_mp3(filename):
    # 初始化pygame库
    pygame.mixer.init()
    
    try:
        # 加载音乐文件
        pygame.mixer.music.load(filename)
        print(f"Playing {filename}")
        
        # 播放音乐
        pygame.mixer.music.play()
        
        # 循环等待音乐播放完毕
        while pygame.mixer.music.get_busy():
            pygame.time.Clock().tick(10)
    except pygame.error as e:
        print(f"Error playing {filename}: {e}")

def play_music(musicname):
    global current_playing
    # 停止当前正在播放的音乐
    if current_playing:
        pygame.mixer.music.stop()
    # 记录当前正在播放的音乐文件名
    current_playing = musicname
    musicname = "tests/mp3/" + musicname
    # 创建并启动一个新线程来播放音乐
    t = threading.Thread(target=play_mp3, args=(musicname,))
    t.start()

    return